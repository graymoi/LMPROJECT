from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_shading(cell, fill_color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), fill_color)
    cell._tc.get_or_add_tcPr().append(shading)

def create_weekly_report():
    doc = Document()
    
    # 设置默认字体
    style = doc.styles['Normal']
    style.font.name = 'SimSun'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    style.font.size = Pt(10.5)
    
    # 标题
    title = doc.add_heading('周工作总结（2.26-3.1）', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 基本信息
    doc.add_paragraph('总结周期：2026年2月26日 - 3月1日')
    doc.add_paragraph('生成时间：2026-03-01')
    doc.add_paragraph('状态：最终版')
    
    doc.add_paragraph('─' * 50)
    
    # 一、总结与计划
    doc.add_heading('一、总结与计划', level=1)
    
    # 1.1 本周工作总结
    doc.add_heading('1.1 本周工作总结', level=2)
    
    # 背街小巷项目
    doc.add_heading('背街小巷系列项目推进', level=3)
    
    table1 = doc.add_table(rows=4, cols=3)
    table1.style = 'Table Grid'
    headers = ['时间', '工作内容', '成果']
    for i, header in enumerate(headers):
        table1.rows[0].cells[i].text = header
        set_cell_shading(table1.rows[0].cells[i], 'D9E2F3')
    
    data1 = [
        ['2.26', '前期调研收尾、找城管委要钱、研究项目包装类型', '明确体检工作需开展、资金来源待确定'],
        ['2.27', '对接城管委规计处魏工、明确政策体系分工', '前期预算500万、与天大联合（导则200万+调研300万）'],
        ['2.28-3.1', '编制招采技术要求', '利用大模型推理能力，完成技术文件编制要求']
    ]
    for i, row_data in enumerate(data1):
        for j, cell_data in enumerate(row_data):
            table1.rows[i+1].cells[j].text = cell_data
    
    doc.add_paragraph()
    doc.add_paragraph('关键进展：')
    doc.add_paragraph('• 政策体系构成明确：导则（天大）+ 调研报告（我们）+ 政策工具库（我们）')
    doc.add_paragraph('• 资金安排明确：市财政只能出资到"体检"之前，实施项目需区里牵头、出资')
    doc.add_paragraph('• 招采文件编制时间节点：周一编制招标任务书，下周三正式汇总完稿')
    
    doc.add_paragraph()
    doc.add_paragraph('困难与问题：')
    doc.add_paragraph('• 王金铭处长拒绝提供各区联系人方式')
    doc.add_paragraph('• 政府"没钱"状态，前期咨询类工作资金解决困难')
    
    # 河西道路项目
    doc.add_heading('河西道路项目申报', level=3)
    
    table2 = doc.add_table(rows=3, cols=3)
    table2.style = 'Table Grid'
    for i, header in enumerate(headers):
        table2.rows[0].cells[i].text = header
        set_cell_shading(table2.rows[0].cells[i], 'D9E2F3')
    
    data2 = [
        ['2.26', '提交发改委', '已申报'],
        ['2.27', '添加附件、讨论绑定逻辑', '发现"小区外道路绑定小区"存在争议']
    ]
    for i, row_data in enumerate(data2):
        for j, cell_data in enumerate(row_data):
            table2.rows[i+1].cells[j].text = cell_data
    
    doc.add_paragraph()
    doc.add_paragraph('后续计划：当作案例继续研究')
    
    # 河东区项目
    doc.add_heading('河东区项目跟踪', level=3)
    
    doc.add_paragraph('项目一：河东老旧小区周边道路改造项目')
    doc.add_paragraph('• 河东城管委年前没赶上报项目')
    doc.add_paragraph('• 周工持续跟踪，下周可能去河东')
    doc.add_paragraph('• 可参考河西项目（P-2026-001）的申报经验')
    
    doc.add_paragraph()
    doc.add_paragraph('项目二：河东区城建谋划项目分类汇总（9个子项目）')
    doc.add_paragraph('• 跟踪负责人：金工')
    doc.add_paragraph('• 部分项目已信息脱敏，作为案例')
    doc.add_paragraph('• 利用政策知识库进行资金拼盘报告输出')
    
    # 政策知识库
    doc.add_heading('政策知识库建设（T-2026-001）', level=3)
    
    table3 = doc.add_table(rows=4, cols=3)
    table3.style = 'Table Grid'
    for i, header in enumerate(headers):
        table3.rows[0].cells[i].text = header
        set_cell_shading(table3.rows[0].cells[i], 'D9E2F3')
    
    data3 = [
        ['2.26', '完善一部分、梳理架构', '知识库架构初步成型'],
        ['2.27', '给同事介绍知识库工具', '明确产物：资金拼盘策略报告、体检报告'],
        ['2.28-3.1', '完善新闻监测Skill、探讨知识管理vs AI应用模式', '形成核心差异文档']
    ]
    for i, row_data in enumerate(data3):
        for j, cell_data in enumerate(row_data):
            table3.rows[i+1].cells[j].text = cell_data
    
    doc.add_paragraph()
    doc.add_paragraph('核心功能：政策咨询、资金申报、项目谋划')
    
    # 个人工具完善
    doc.add_heading('个人工具完善', level=3)
    
    table4 = doc.add_table(rows=4, cols=3)
    table4.style = 'Table Grid'
    headers2 = ['时间', '工作内容', '状态']
    for i, header in enumerate(headers2):
        table4.rows[0].cells[i].text = header
        set_cell_shading(table4.rows[0].cells[i], 'D9E2F3')
    
    data4 = [
        ['2.26', '需要AI读取Excel表格对比数据', '✅ 已解决'],
        ['2.26', '需要AI帮忙归档河西区项目', '✅ 已解决'],
        ['2.28-3.1', '建立对话日志存档规则', '✅ 已完成']
    ]
    for i, row_data in enumerate(data4):
        for j, cell_data in enumerate(row_data):
            table4.rows[i+1].cells[j].text = cell_data
    
    # 1.2 下周工作计划
    doc.add_heading('1.2 下周工作计划', level=2)
    
    table5 = doc.add_table(rows=7, cols=4)
    table5.style = 'Table Grid'
    headers3 = ['优先级', '工作内容', '时间节点', '负责人']
    for i, header in enumerate(headers3):
        table5.rows[0].cells[i].text = header
        set_cell_shading(table5.rows[0].cells[i], 'D9E2F3')
    
    data5 = [
        ['高', '编制招标任务书、天大做预算', '周一', '我们+天大'],
        ['高', '正式汇总完稿', '下周三', '我们+天大'],
        ['高', '继续推进背街小巷政策体系编制', '全周', '我们'],
        ['中', '河西区项目案例研究', '全周', '我们'],
        ['中', '河东区项目跟踪', '全周', '周工'],
        ['中', '知识库持续完善', '全周', '我们']
    ]
    for i, row_data in enumerate(data5):
        for j, cell_data in enumerate(row_data):
            table5.rows[i+1].cells[j].text = cell_data
    
    # 二、新掌握的知识
    doc.add_heading('二、新掌握的知识、常识列表', level=1)
    
    doc.add_heading('2.1 政策知识', level=2)
    doc.add_paragraph('• 中央预算内投资 vs 中央财政补助资金：两个概念的区别需掌握')
    doc.add_paragraph('• 前期策划精准谋划取费依据：需要学习掌握')
    doc.add_paragraph('• 背街小巷概念范畴：支路、支路以下、无名路；更像一场行动/运动')
    doc.add_paragraph('• 政府资金运作模式：市财政只能出资到"体检"之前；实施项目需区里牵头、出资')
    
    doc.add_heading('2.2 方法论知识', level=2)
    doc.add_paragraph('• 项目资金拼盘策略：产物之一：资金拼盘策略报告')
    doc.add_paragraph('• 体检报告编制：产物之二：体检报告')
    doc.add_paragraph('• 知识管理 vs AI应用模式：本地化知识管理的核心优势')
    
    doc.add_heading('2.3 业务认知升级', level=2)
    doc.add_paragraph('"推进项目，现在需要传统设计师去推进，从找到钱，找到政策资金、再找到政策资金的跑通的流程"')
    doc.add_paragraph('核心洞察：')
    doc.add_paragraph('• 政府"没钱"状态是常态')
    doc.add_paragraph('• 前期咨询类工作的资金解决是服务方难题')
    doc.add_paragraph('• 项目成型前，大家都希望投入能有保证')
    
    # 三、适合给大家科普的附件
    doc.add_heading('三、适合给大家科普的附件', level=1)
    
    doc.add_heading('3.1 新闻监测记录', level=2)
    doc.add_paragraph('• 2026年第一批936亿元超长期特别国债支持设备更新资金已下达')
    doc.add_paragraph('• 住建部、自然资源部发布2026城市更新新政')
    doc.add_paragraph('• 2026年提前批"两重"项目清单和中央预算内投资已下达（约2950亿元）')
    
    doc.add_heading('3.2 阶段成果A4报告', level=2)
    doc.add_paragraph('• P-2026-002：背街小巷精细化治理_A4介绍')
    doc.add_paragraph('• P-2026-001：河西道路项目_A4介绍')
    doc.add_paragraph('• T-2026-001：精准谋划知识库_A4介绍')
    
    # 四、复用的数据成果
    doc.add_heading('四、复用的数据成果', level=1)
    
    doc.add_heading('4.1 知识库架构', level=2)
    doc.add_paragraph('精准谋划知识库结构：')
    doc.add_paragraph('• 00_政策法规/（政策咨询核心）')
    doc.add_paragraph('• 01_资金申报指南/（资金申报核心）')
    doc.add_paragraph('• 02_项目案例库/（精准谋划参考）')
    doc.add_paragraph('• 03_谋划方法论/（精准谋划核心）')
    doc.add_paragraph('• 04_专家问答库/（辅助推理）')
    
    doc.add_heading('4.2 Skills', level=2)
    doc.add_paragraph('• 新闻监测：自动监测新闻和政策')
    doc.add_paragraph('• 日工作记录：每日碎片化记录+知识萃取')
    doc.add_paragraph('• 周工作总结：周总结生成+知识汇总')
    
    doc.add_heading('4.3 规则库', level=2)
    doc.add_paragraph('• 对话日志存档规则：每10轮对话自动存档')
    doc.add_paragraph('• 文件命名规则：日总结、周总结、新闻监测等')
    doc.add_paragraph('• 输入输出分离规则：用户输入AI只读、AI输出用户可编辑')
    
    # 五、个人工具层面
    doc.add_heading('五、个人工具层面', level=1)
    
    doc.add_heading('5.1 整体进度', level=2)
    doc.add_paragraph('已完成：')
    doc.add_paragraph('• 工作总结工作流设计')
    doc.add_paragraph('• 文件夹结构搭建')
    doc.add_paragraph('• 输入输出分离机制')
    doc.add_paragraph('• 对话日志存档规则')
    doc.add_paragraph('• 新闻监测Skill')
    doc.add_paragraph('• 河西区项目归档（P-2026-001）')
    doc.add_paragraph('• 背街小巷项目归档（P-2026-002）')
    doc.add_paragraph('• 精准谋划知识库归档（T-2026-001）')
    doc.add_paragraph('• AI读取Excel表格对比数据')
    doc.add_paragraph('• 三个项目A4报告')
    
    doc.add_paragraph()
    doc.add_paragraph('进行中：')
    doc.add_paragraph('• 政策知识库完善')
    doc.add_paragraph('• 项目库建设')
    doc.add_paragraph('• 个人知识库建设')
    
    doc.add_heading('5.2 从项目库萃取的内容', level=2)
    doc.add_paragraph('• 背街小巷系列项目（P-2026-002）：对接人员、资金安排、工作分工、"六个一"任务、四级联动机制')
    doc.add_paragraph('• 老旧小区周边道路改造项目（P-2026-001）：项目基本信息、谋划过程、关键经验、筛选标准、攻防问答要点')
    doc.add_paragraph('• 天津城管委组织架构（T-2026-002）：28个处室、14个下属单位、职责分工')
    doc.add_paragraph('• 精准谋划知识库（T-2026-001）：5层架构、核心功能、成果文件、应用场景')
    
    # 保存文档
    doc.save(r'd:\LMAI\000打杂工具箱\输出成果\周工作总结\2.26-3.1周工作总结.docx')
    print("Word文档已生成：2.26-3.1周工作总结.docx")

if __name__ == '__main__':
    create_weekly_report()
