from docx import Document
from pathlib import Path

# Read the meeting minutes document
doc_path = Path(r'e:\李萌\26-发改委谋划天津项目\260311\会议纪要0311Z.docx')
doc = Document(doc_path)

# Extract all text content
full_text = []
for paragraph in doc.paragraphs:
    if paragraph.text.strip():
        full_text.append(paragraph.text)

# Extract tables
for table in doc.tables:
    full_text.append("\n【表格内容】")
    for row in table.rows:
        row_text = []
        for cell in row.cells:
            if cell.text.strip():
                row_text.append(cell.text.strip())
        if row_text:
            full_text.append(" | ".join(row_text))

# Print the content
print("=== 会议纪要内容 ===")
for i, text in enumerate(full_text, 1):
    print(f"{i}. {text}")
