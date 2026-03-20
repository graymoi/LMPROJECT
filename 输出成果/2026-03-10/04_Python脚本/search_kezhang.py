from docx import Document
from pathlib import Path

# Read the meeting minutes document
doc_path = Path(r'e:\李萌\26-发改委谋划天津项目\260311\会议纪要0311Z.docx')
doc = Document(doc_path)

# Extract all text content
full_text = []
for i, paragraph in enumerate(doc.paragraphs, 1):
    if paragraph.text.strip():
        full_text.append(paragraph.text.strip())

# Extract tables
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            if cell.text.strip():
                full_text.append(cell.text.strip())

# Search for "科长"
print("=== 搜索'科长' ===")
found = False
for i, text in enumerate(full_text, 1):
    if '科长' in text:
        print(f"第{i}段: {text}")
        found = True

if not found:
    print("未找到'科长'")

# Search for "郑"
print("\n=== 搜索'郑' ===")
found = False
for i, text in enumerate(full_text, 1):
    if '郑' in text:
        print(f"第{i}段: {text}")
        found = True

if not found:
    print("未找到'郑'")

# Print all text for reference
print("\n=== 完整文本 ===")
for i, text in enumerate(full_text, 1):
    print(f"{i}. {text}")
