from docx import Document
from pathlib import Path
import re

# Read the meeting minutes document
doc_path = Path(r'e:\李萌\26-发改委谋划天津项目\260311\会议纪要0311Z.docx')
doc = Document(doc_path)

# Extract all text content with paragraph numbers
full_text = []
for i, paragraph in enumerate(doc.paragraphs, 1):
    if paragraph.text.strip():
        full_text.append(f"{i}. {paragraph.text}")

# Extract tables with table numbers
for table_num, table in enumerate(doc.tables, 1):
    full_text.append(f"\n【表格 {table_num}】")
    for row_num, row in enumerate(table.rows, 1):
        row_text = []
        for cell in row.cells:
            if cell.text.strip():
                row_text.append(cell.text.strip())
        if row_text:
            full_text.append(f"  行{row_num}: {' | '.join(row_text)}")

# Print content
print("=== 会议纪要详细内容 ===")
for text in full_text:
    print(text)
