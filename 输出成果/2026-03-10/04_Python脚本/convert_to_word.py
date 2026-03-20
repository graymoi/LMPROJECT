from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from pathlib import Path

# Create a new Document
doc = Document()

# Set page orientation to landscape
section = doc.sections[0]
section.page_height = Inches(8.27)  # A4 height
section.page_width = Inches(11.69)   # A4 width (landscape)
section.left_margin = Inches(0.6)
section.right_margin = Inches(0.6)
section.top_margin = Inches(0.6)
section.bottom_margin = Inches(0.6)

# Function to add heading
def add_heading(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if level == 1:
        run = heading.runs[0]
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(26, 84, 144)
        run.font.bold = True
    elif level == 2:
        run = heading.runs[0]
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(44, 123, 182)
        run.font.bold = True
    elif level == 3:
        run = heading.runs[0]
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(68, 68, 68)
        run.font.bold = True

# Function to add paragraph
def add_paragraph(doc, text, bold=False):
    p = doc.add_paragraph(text)
    run = p.runs[0]
    run.font.size = Pt(10)
    run.font.bold = bold
    return p

# Function to add table
def add_table(doc, data, col_widths=None):
    table = doc.add_table(rows=len(data), cols=len(data[0]))
    table.style = 'Light Grid Accent 1'

    for i, row_data in enumerate(data):
        row = table.rows[i]
        for j, cell_data in enumerate(row_data):
            cell = row.cells[j]
            cell.text = str(cell_data)

            # Set font size
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)

            # Set column widths if provided
            if col_widths and i == 0:
                cell.width = Inches(col_widths[j])

            # Style header row
            if i == 0:
                cell.background_color = RGBColor(240, 244, 248)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.size = Pt(10)

    return table

# Header info
header_table = doc.add_table(rows=1, cols=3)
header_table.style = 'Table Grid'
header_cells = header_table.rows[0].cells
header_cells[0].text = '汇报单位：天津市河西区城市管理委员会'
header_cells[1].text = '汇报日期：2026年3月10日'
header_cells[2].text = '项目编号：P-2026-001'

for cell in header_cells:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.size = Pt(9)

doc.add_paragraph().add_run().add_break()

# Section 1: 项目概况
add_heading(doc, '一、项目概况')

overview_data = [
    ['项目名称', '河西区老旧小区周边衔接道路基础设施配套项目'],
    ['项目类型', '城市更新-基础设施改造'],
    ['申报状态', '已申报，等待资金批复'],
    ['道路数量', '53条（20条+33条）'],
    ['改造长度', '36.69 km'],
    ['改造面积', '64.45万m²'],
    ['总投资', '19151.59万元'],
    ['工程费用', '15382.06万元'],
    ['拟申请中央资金', '15321.27万元']
]

overview_table = doc.add_table(rows=len(overview_data), cols=2)
overview_table.style = 'Light Grid Accent 1'

for i, row_data in enumerate(overview_data):
    row = overview_table.rows[i]
    for j, cell_data in enumerate(row_data):
        cell = row.cells[j]
        cell.text = str(cell_data)

        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(10)

        # Style first column
        if j == 0:
            cell.background_color = RGBColor(240, 244, 248)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True

doc.add_paragraph().add_run().add_break()

# Section 2: 工作周期
add_heading(doc, '二、工作周期')
add_paragraph(doc, '项目周期：23天（2026年2月5日-2月27日）', bold=True)

timeline_text = '''
2月5日      2月6日      2月10日      2月13日      2月25日      2月27日
  │           │           │            │            │            │
项目启动    部门对接    方案编制    项目分拆    要件办理    申报提交
  │           │           │            │            │            │
明确方向    数据收集    方案设计    20+33分拆    前置要件    完成申报
  │           │           └────────────┴────────────┘
  │           │                   │
  │           │             2.10-2.25 为图纸、实施方案、资金申请报告主要周期
'''

add_paragraph(doc, timeline_text)

doc.add_paragraph().add_run().add_break()

# Section 3: 多方协调工作
add_heading(doc, '三、多方协调工作')
add_paragraph(doc, '协调时间：2026年2月4日-2月28日（25天）', bold=True)
add_paragraph(doc, '协调单位：区发改委、区住建委、区城管委、市发改委', bold=True)
doc.add_paragraph().add_run().add_break()

coord_data = [
    ['时间', '协调单位', '协调内容', '工作成果'],
    ['2月4日', '项目前期准备', '在市发改委召开座谈会', '明确项目方向'],
    ['2月5日', '区发改委、住建委、城管委', '河西区发改委对接会议', '探讨项目申报方向，协调数据共享机制'],
    ['2月6日上午', '区住建委', '住建委座谈', '收集数据资料，排查城更范围'],
    ['2月6日下午', '区城管委', '河西区城管委调取数据', '拆分区管道路中的支路，调取背街小巷统计数据'],
    ['2月9日上午', '区住建委', '从河西区住建委取得城更台账以及21-25年老旧小区改造清单', '取得城更既有项目以及谋划储备项目的台账，以及21-25年老旧小区改造清单'],
    ['2月9日下午', '市发改委', '向市发改委汇报', '汇报整体进度以及项目拟申报资金'],
    ['2月10日', '方案编制启动', '承担项目全程的设计及资金申报工作', '《天津市河西区老旧小区周边衔接道路基础设施配套项目》前期资料包'],
    ['2月11-24日', '方案编制', '正式编制实施方案及资金申请报告', '以资金申请报告、实施方案为主要技术成果'],
    ['2月25-26日', '区住建委、市发改委', '前置要件办理', '前往河西区住建委、市发改委，协助办理各种前置要件'],
    ['2月26日', '成果完成', '完成所有成果', '天津市河西区老旧小区周边衔接道路基础设施配套项目完整成果包'],
    ['2月27日', '申报提交', '完成申报提交', '']
]

coord_table = doc.add_table(rows=len(coord_data), cols=4)
coord_table.style = 'Light Grid Accent 1'

for i, row_data in enumerate(coord_data):
    row = coord_table.rows[i]
    for j, cell_data in enumerate(row_data):
        cell = row.cells[j]
        cell.text = str(cell_data)

        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(8)

        # Style header row
        if i == 0:
            cell.background_color = RGBColor(240, 244, 248)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(9)

doc.add_paragraph().add_run().add_break()

add_paragraph(doc, '协调工作特点：', bold=True)
add_paragraph(doc, '✅ 跨部门协调：市区城管委、市发改委、区发改委、区住建委多部门联动')
add_paragraph(doc, '✅ 高频次对接：25天内完成9次重要协调对接')
add_paragraph(doc, '✅ 全流程覆盖：从项目启动到申报提交，全流程协调')
add_paragraph(doc, '✅ 问题导向：针对每个关键节点进行针对性协调')

doc.add_paragraph().add_run().add_break()

# Section 4: 重要成果
add_heading(doc, '四、重要成果')

# 4.1 技术成果
add_heading(doc, '4.1 技术成果', level=2)

tech_data = [
    ['成果类型', '成果内容', '数量'],
    ['实施方案', '项目实施方案（P-2026-001-A/B）', '2份'],
    ['资金申请报告', '资金申请报告（P-2026-001-A/B）', '2份'],
    ['道路清单', '53条道路详细清单', '1份'],
    ['综合估算表', '工程综合估算表', '2份'],
    ['绩效目标表', '项目绩效目标表', '2份']
]

tech_table = doc.add_table(rows=len(tech_data), cols=3)
tech_table.style = 'Light Grid Accent 1'

for i, row_data in enumerate(tech_data):
    row = tech_table.rows[i]
    for j, cell_data in enumerate(row_data):
        cell = row.cells[j]
        cell.text = str(cell_data)

        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(9)

        # Style header row
        if i == 0:
            cell.background_color = RGBColor(240, 244, 248)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(10)

doc.add_paragraph().add_run().add_break()

# 4.2 文件成果
add_heading(doc, '4.2 文件成果', level=2)

file_data = [
    ['文件类别', '文件数量', '说明'],
    ['申报表', '2份', '项目申报表'],
    ['绩效目标', '2份', '项目绩效目标'],
    ['前期资料', '8份', '立项、批复、说明等'],
    ['资金材料', '2份', '财政承诺文件'],
    ['技术文件', '6份', '资金申请报告、实施方案等'],
    ['其他材料', '8份', '纳入重大项目库、信用报告等'],
    ['软建设', '5份', '政策支撑文件'],
    ['合计', '33份', 'P-2026-001-A/B共42份文件']
]

file_table = doc.add_table(rows=len(file_data), cols=3)
file_table.style = 'Light Grid Accent 1'

for i, row_data in enumerate(file_data):
    row = file_table.rows[i]
    for j, cell_data in enumerate(row_data):
        cell = row.cells[j]
        cell.text = str(cell_data)

        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(9)

        # Style header row
        if i == 0:
            cell.background_color = RGBColor(240, 244, 248)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(10)

        # Style total row
        if i == len(file_data) - 1:
            cell.background_color = RGBColor(255, 249, 230)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True

doc.add_paragraph().add_run().add_break()

# 4.3 管理成果
add_heading(doc, '4.3 管理成果', level=2)
add_paragraph(doc, '✅ 项目谋划方法论：精准卡位、快速响应、避免重复、系统思维')
add_paragraph(doc, '✅ 攻防问答要点：9个常见评审质疑与应对策略')
add_paragraph(doc, '✅ 硬支撑清单：地图+实景图、无障碍设施列表、叠加老小区范围图、道路体检机制')

doc.add_paragraph().add_run().add_break()

# Section 5: 下一步工作计划
add_heading(doc, '五、下一步工作计划')

add_heading(doc, '近期工作（1-2周）', level=3)
add_paragraph(doc, '• 跟踪中央预算内资金批复进度')
add_paragraph(doc, '• 准备合同签署材料')
add_paragraph(doc, '• 施工单位招标准备')

add_heading(doc, '中期工作（1-2月）', level=3)
add_paragraph(doc, '• 资金到位后立即启动施工')
add_paragraph(doc, '• 按照项目段分批实施')
add_paragraph(doc, '• 建立项目管理机制')

add_heading(doc, '长期工作（3-6月）', level=3)
add_paragraph(doc, '• 项目完成后组织验收')
add_paragraph(doc, '• 总结项目经验')
add_paragraph(doc, '• 完成项目归档')

doc.add_paragraph().add_run().add_break()

# Section 6: 合同签署
add_heading(doc, '六、合同签署')
add_paragraph(doc, '建议签署时间：立即签署合同，积极申请中央资金', bold=True)
doc.add_paragraph().add_run().add_break()
add_paragraph(doc, '合同主要内容：', bold=True)
add_paragraph(doc, '• 项目范围：53条道路改造')
add_paragraph(doc, '• 项目投资：19151.59万元')
add_paragraph(doc, '• 拟申请中央资金：15321.27万元')
add_paragraph(doc, '• 项目周期：20个月（2026年5月-2027年12月）')
add_paragraph(doc, '• 质量标准：符合国家相关标准')

doc.add_paragraph().add_run().add_break()

# Footer
footer_table = doc.add_table(rows=1, cols=3)
footer_table.style = 'Table Grid'
footer_cells = footer_table.rows[0].cells
footer_cells[0].text = '汇报日期: 2026-03-10'
footer_cells[1].text = ''
footer_cells[2].text = '汇报单位: 建研院'

for cell in footer_cells:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(128, 128, 128)

# Save the document
output_path = Path(r'e:\李萌\输出成果\2026-03-10\P-2026-001_河西区老旧小区周边道路改造项目工作汇报_横板.docx')
doc.save(output_path)

print(f"Word document generated successfully: {output_path}")
