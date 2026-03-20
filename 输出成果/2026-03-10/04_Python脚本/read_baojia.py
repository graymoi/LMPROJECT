from docx import Document
from pathlib import Path

# Read the document
doc_path = Path(r'e:\李萌\26-发改委谋划天津项目\260311\中国建筑科学研究院   报价说明.docx')
doc = Document(doc_path)

# Extract all text content
print("=== 文档内容 ===\n")

full_text = []
for i, paragraph in enumerate(doc.paragraphs, 1):
    if paragraph.text.strip():
        full_text.append(paragraph.text.strip())
        print(f"{i}. {paragraph.text}")

# Extract tables
print("\n=== 表格内容 ===\n")
for table_num, table in enumerate(doc.tables, 1):
    print(f"\n【表格 {table_num}】")
    for row_num, row in enumerate(table.rows, 1):
        row_text = []
        for cell in row.cells:
            if cell.text.strip():
                row_text.append(cell.text.strip())
        if row_text:
            print(f"  行{row_num}: {' | '.join(row_text)}")
